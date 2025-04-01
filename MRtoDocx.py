from docx import Document
from docx.shared import Inches
from wand.image import Image
import os


# def MRtoDocx(workbook_data, output_path):
#     for keys in workbook_data:
#         document = Document()
#         document.add_paragraph(f"文件名是{keys}；")
#         document.add_paragraph(f"工艺标准标题是{workbook_data[keys].get('B2', '')}；")
#         document.add_paragraph(f"工艺标准编号是{workbook_data[keys].get('A2', '')}；")
#         document.add_paragraph(f"工艺标准内容详细描述是{workbook_data[keys].get('B4', '')}；")
#         document.add_paragraph(f"工艺标准要求是{workbook_data[keys].get('B5', '')}；")
#         image_path = workbook_data[keys].get('关联图片A4')
#         if image_path:
#             prefix,extension=os.path.splitext(workbook_data[keys]['关联图片A4'])
#             if extension != 'png':
#                 png_path = prefix+'.png'
#                 try:
#                     with Image(filename=workbook_data[keys]['关联图片A4']) as img:
#                         img.format = 'png'
#                         img.save(filename=png_path)
#                     image_path = png_path
#                 except Exception as e:
#                     print(f"转换图片格式时出错：{e}")
#             try:
#                 document.add_picture(image_path,width=Inches(4))
#             except Exception as e:
#                 print(f"{image_path}图片不存在或无法添加：{e}")
#         key=keys.replace('.','')
#         kb_path = os.path.join(output_path, f"{key}.docx")
#         document.save(kb_path)

def MRtoDocx(workbook_data, output_path):
    document = Document()
    for keys in workbook_data:
        document.add_paragraph(f"文件名是{keys}；")
        document.add_paragraph(f"工艺标准标题是{workbook_data[keys].get('B2', '')}；")
        document.add_paragraph(f"工艺标准编号是{workbook_data[keys].get('A2', '')}；")
        document.add_paragraph(f"工艺标准内容详细描述是{workbook_data[keys].get('B4', '')}；")
        document.add_paragraph(f"工艺标准要求是{workbook_data[keys].get('B5', '')}；")
        image_path = workbook_data[keys].get('关联图片A4')
        if image_path:
            prefix,extension=os.path.splitext(workbook_data[keys]['关联图片A4'])
            if extension != 'png':
                png_path = prefix+'.png'
                try:
                    with Image(filename=workbook_data[keys]['关联图片A4']) as img:
                        img.format = 'png'
                        img.save(filename=png_path)
                    image_path = png_path
                except Exception as e:
                    print(f"转换图片格式时出错：{e}")
            try:
                document.add_picture(image_path,width=Inches(4))
            except Exception as e:
                print(f"{image_path}图片不存在或无法添加：{e}")
        document.add_paragraph("!")
    kb_path = os.path.join(output_path, "konwledgebase.docx")
    document.save(kb_path)

